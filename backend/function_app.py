import azure.functions as func
import logging
import os
import re
import json
from io import BytesIO
from datetime import datetime, timezone, timedelta

from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeDocumentRequest, DocumentContentFormat, DocumentAnalysisFeature
from azure.core.credentials import AzureKeyCredential

from docx import Document
from bs4 import BeautifulSoup

from openai import AzureOpenAI

from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions

app = func.FunctionApp()

# --- Tabela de precos (USD). Ajuste aqui se a fatura mostrar valores diferentes. ---
CUSTO_INPUT_POR_1M = 0.25
CUSTO_OUTPUT_POR_1M = 2.00
CUSTO_POR_PAGINA_DI = 0.01

CONTAINER_ENTRADA = "entrada"
CONTAINER_SAIDA = "saida"


# ============================================================
# Clientes
# ============================================================

def get_docint_client() -> DocumentIntelligenceClient:
    endpoint = os.environ["DOCINT_ENDPOINT"].rstrip("/")
    key = os.environ["DOCINT_KEY"]
    return DocumentIntelligenceClient(endpoint=endpoint, credential=AzureKeyCredential(key))


def get_openai_client() -> AzureOpenAI:
    return AzureOpenAI(
        azure_endpoint=os.environ["OPENAI_ENDPOINT"].rstrip("/"),
        api_key=os.environ["OPENAI_KEY"],
        api_version="2024-10-21",
    )


def get_blob_service() -> BlobServiceClient:
    return BlobServiceClient.from_connection_string(os.environ["PDF_STORAGE"])


# ============================================================
# Traducao
# ============================================================

def traduzir_texto(conteudo: str):
    client = get_openai_client()
    deployment = os.environ["OPENAI_DEPLOYMENT"]

    system_prompt = (
        "You are a professional translator. Translate the following document content "
        "into English. Preserve ALL markdown formatting (headings with #, etc.) and keep "
        "any HTML tables (<table>, <tr>, <td>) exactly as structured, translating only the "
        "text inside them. Do not add explanations or comments; return only the translated content."
    )

    resposta = client.chat.completions.create(
        model=deployment,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": conteudo},
        ],
    )
    return resposta.choices[0].message.content, resposta.usage


# ============================================================
# Montagem de DOCX (devolve bytes, nao salva em disco)
# ============================================================

def adicionar_tabela_docx(doc, html_tabela):
    soup = BeautifulSoup(html_tabela, "html.parser")
    linhas = soup.find_all("tr")
    if not linhas:
        return

    num_colunas = len(linhas[0].find_all(["td", "th"]))
    if num_colunas == 0:
        return

    tabela = doc.add_table(rows=0, cols=num_colunas)
    tabela.style = "Table Grid"

    for tr in linhas:
        celulas = tr.find_all(["td", "th"])
        linha_docx = tabela.add_row().cells
        for i, celula in enumerate(celulas):
            if i < num_colunas:
                linha_docx[i].text = celula.get_text(strip=True)


def montar_docx_bytes(conteudo) -> bytes:
    doc = Document()
    # Remove os comentarios de layout que o Document Intelligence insere
    # (ex.: <!-- PageBreak -->, <!-- PageNumber="2." -->) pra nao virarem texto cru.
    conteudo = re.sub(r"<!--.*?-->", "", conteudo, flags=re.DOTALL)
    partes = re.split(r"(<table>.*?</table>)", conteudo, flags=re.DOTALL)

    for parte in partes:
        parte = parte.strip()
        if not parte:
            continue

        if parte.startswith("<table>"):
            adicionar_tabela_docx(doc, parte)
        else:
            for linha in parte.split("\n"):
                linha = linha.strip()
                if not linha:
                    continue
                match = re.match(r"^(#{1,6})\s+(.*)", linha)
                if match:
                    nivel = len(match.group(1))
                    doc.add_heading(match.group(2), level=nivel)
                else:
                    doc.add_paragraph(linha)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ============================================================
# Envio pro container de saida
# ============================================================

def enviar_para_saida(nome_base, docx_fiel, docx_traduzido, metadata_dict):
    blob_service = get_blob_service()
    container = blob_service.get_container_client(CONTAINER_SAIDA)

    arquivos = {
        f"{nome_base}/fiel.docx": docx_fiel,
        f"{nome_base}/traduzido.docx": docx_traduzido,
        f"{nome_base}/metadata.json": json.dumps(metadata_dict, indent=2, ensure_ascii=False).encode("utf-8"),
    }

    for caminho, dados in arquivos.items():
        container.upload_blob(name=caminho, data=dados, overwrite=True)
        logging.info(f"[PIPELINE] Enviado pro Blob: {CONTAINER_SAIDA}/{caminho}")


# ============================================================
# Trigger principal — processa o PDF que cai no container "entrada"
# ============================================================

@app.blob_trigger(
    arg_name="pdf",
    path="entrada/{name}",
    connection="PDF_STORAGE"
)
def processar_pdf(pdf: func.InputStream):
    inicio = datetime.now(timezone.utc)

    logging.info(f"[PIPELINE] PDF recebido: {pdf.name}")
    logging.info(f"[PIPELINE] Tamanho: {pdf.length} bytes")

    pdf_bytes = pdf.read()
    nome_base = os.path.splitext(os.path.basename(pdf.name))[0]

    # --- Fase 2: extracao com Document Intelligence ---
    client = get_docint_client()
    logging.info("[PIPELINE] Enviando pro Document Intelligence (prebuilt-layout)...")
    poller = client.begin_analyze_document(
        "prebuilt-layout",
        AnalyzeDocumentRequest(bytes_source=pdf_bytes),
        output_content_format=DocumentContentFormat.MARKDOWN,
        # Add-on de alta resolucao: melhora leitura de texto denso e simbolos ambiguos.
        features=[DocumentAnalysisFeature.OCR_HIGH_RESOLUTION],
    )
    result = poller.result()

    num_paginas = len(result.pages) if result.pages else 0
    num_tabelas = len(result.tables) if result.tables else 0
    logging.info(f"[PIPELINE] Extracao concluida: {num_paginas} paginas, {num_tabelas} tabelas")

    conteudo = result.content

    # --- Fase 3: docx fiel (em memoria) ---
    docx_fiel = montar_docx_bytes(conteudo)
    logging.info("[PIPELINE] DOCX fiel montado.")

    # --- Fase 4: traducao + docx traduzido (em memoria) ---
    logging.info("[PIPELINE] Traduzindo com Azure OpenAI (gpt-5-mini)...")
    conteudo_traduzido, usage = traduzir_texto(conteudo)
    docx_traduzido = montar_docx_bytes(conteudo_traduzido)
    logging.info("[PIPELINE] DOCX traduzido montado.")

    # --- Fase 5: metadados, tokens e custos ---
    fim = datetime.now(timezone.utc)
    duracao_seg = (fim - inicio).total_seconds()

    tokens_input = usage.prompt_tokens
    tokens_output = usage.completion_tokens
    tokens_total = usage.total_tokens

    custo_di = num_paginas * CUSTO_POR_PAGINA_DI
    custo_openai_input = (tokens_input / 1_000_000) * CUSTO_INPUT_POR_1M
    custo_openai_output = (tokens_output / 1_000_000) * CUSTO_OUTPUT_POR_1M
    custo_openai = custo_openai_input + custo_openai_output
    custo_total = custo_di + custo_openai

    metadata = {
        "arquivo": {
            "nome": os.path.basename(pdf.name),
            "tamanho_bytes": pdf.length,
        },
        "processamento": {
            "inicio_utc": inicio.isoformat(),
            "fim_utc": fim.isoformat(),
            "duracao_segundos": round(duracao_seg, 2),
        },
        "extracao": {
            "paginas": num_paginas,
            "tabelas": num_tabelas,
            "caracteres_extraidos": len(conteudo),
        },
        "traducao": {
            "modelo": os.environ["OPENAI_DEPLOYMENT"],
            "idioma_destino": "en",
            "tokens_input": tokens_input,
            "tokens_output": tokens_output,
            "tokens_total": tokens_total,
        },
        "custos_estimados_usd": {
            "document_intelligence": round(custo_di, 6),
            "openai_input": round(custo_openai_input, 6),
            "openai_output": round(custo_openai_output, 6),
            "openai_total": round(custo_openai, 6),
            "total": round(custo_total, 6),
        },
        "saidas": {
            "docx_fiel": f"{nome_base}/fiel.docx",
            "docx_traduzido": f"{nome_base}/traduzido.docx",
            "metadata": f"{nome_base}/metadata.json",
        },
    }

    # --- Persistencia: envia as 3 saidas pro container "saida" ---
    enviar_para_saida(nome_base, docx_fiel, docx_traduzido, metadata)

    logging.info(f"[PIPELINE] Concluido. Custo total: ${custo_total:.6f} USD | Tokens: {tokens_total}")


# ============================================================
# Endpoint HTTP: lista os documentos processados (pro frontend)
# ============================================================

@app.route(route="status", auth_level=func.AuthLevel.ANONYMOUS)
def status(req: func.HttpRequest) -> func.HttpResponse:
    blob_service = get_blob_service()
    container = blob_service.get_container_client(CONTAINER_SAIDA)

    account_key = blob_service.credential.account_key
    account_name = blob_service.account_name

    documentos = {}

    # Lista todos os blobs do container e agrupa por "pasta" (nome do documento).
    for blob in container.list_blobs():
        partes = blob.name.split("/", 1)
        if len(partes) != 2:
            continue
        nome_doc, arquivo = partes

        if nome_doc not in documentos:
            documentos[nome_doc] = {"nome": nome_doc, "arquivos": {}, "metadata": None}

        # SAS de leitura valido por 1 hora pra esse arquivo.
        sas = generate_blob_sas(
            account_name=account_name,
            container_name=CONTAINER_SAIDA,
            blob_name=blob.name,
            account_key=account_key,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.now(timezone.utc) + timedelta(hours=1),
        )
        url = f"https://{account_name}.blob.core.windows.net/{CONTAINER_SAIDA}/{blob.name}?{sas}"
        documentos[nome_doc]["arquivos"][arquivo] = url

    # Pra cada documento, baixa e embute o conteudo do metadata.json.
    for nome_doc, doc in documentos.items():
        try:
            blob_client = container.get_blob_client(f"{nome_doc}/metadata.json")
            conteudo = blob_client.download_blob().readall()
            doc["metadata"] = json.loads(conteudo)
        except Exception:
            doc["metadata"] = None

    # Ordena por inicio do processamento (mais recente primeiro).
    lista = list(documentos.values())
    lista.sort(
        key=lambda d: d["metadata"]["processamento"]["inicio_utc"] if d.get("metadata") else "",
        reverse=True,
    )

    return func.HttpResponse(
        json.dumps({"documentos": lista}, ensure_ascii=False),
        mimetype="application/json",
        headers={
            "Access-Control-Allow-Origin": "*",
            "Cache-Control": "no-cache, no-store, must-revalidate",
        },
    )


# ============================================================
# Endpoint HTTP: gera um SAS de escrita pro frontend subir o PDF
# ============================================================

@app.route(route="upload-url", auth_level=func.AuthLevel.ANONYMOUS)
def upload_url(req: func.HttpRequest) -> func.HttpResponse:
    filename = req.params.get("filename")
    if not filename:
        return func.HttpResponse(
            json.dumps({"erro": "Informe o parametro filename"}),
            status_code=400,
            mimetype="application/json",
            headers={"Access-Control-Allow-Origin": "*"},
        )

    # So aceita PDF (o pipeline so processa PDF).
    if not filename.lower().endswith(".pdf"):
        return func.HttpResponse(
            json.dumps({"erro": "Somente arquivos .pdf sao aceitos"}),
            status_code=400,
            mimetype="application/json",
            headers={"Access-Control-Allow-Origin": "*"},
        )

    blob_service = get_blob_service()
    account_key = blob_service.credential.account_key
    account_name = blob_service.account_name

    # SAS de ESCRITA valido por 10 minutos, limitado a esse arquivo no "entrada".
    sas = generate_blob_sas(
        account_name=account_name,
        container_name=CONTAINER_ENTRADA,
        blob_name=filename,
        account_key=account_key,
        permission=BlobSasPermissions(write=True, create=True),
        expiry=datetime.now(timezone.utc) + timedelta(minutes=10),
    )
    url = f"https://{account_name}.blob.core.windows.net/{CONTAINER_ENTRADA}/{filename}?{sas}"

    return func.HttpResponse(
        json.dumps({"uploadUrl": url}),
        mimetype="application/json",
        headers={
            "Access-Control-Allow-Origin": "*",
            "Cache-Control": "no-store",
        },
    )